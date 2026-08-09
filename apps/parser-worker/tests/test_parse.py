import sys
from pathlib import Path
from uuid import uuid4

import ezdxf
from docx import Document
from fastapi.testclient import TestClient
from openpyxl import Workbook
from pytest import MonkeyPatch

from app import main as main_module
from app.config import Settings
from app.main import create_app
from app.parsers import image as image_parser
from app.schemas import ParsedElement, PreviewArtifact

TOKEN = "test-internal-token"  # noqa: S105 -- non-secret test fixture


def make_fake_dwg_converter(path: Path) -> None:
    path.write_text(
        f"""#!{sys.executable}
import sys
from pathlib import Path

import ezdxf

output_dir = Path(sys.argv[2])
source_dir = Path(sys.argv[1])
if sys.argv[3:] != ["ACAD2018", "DXF", "0", "1", "*.dwg"]:
    raise SystemExit("unexpected ODA command schema")
source_name = next(source_dir.glob("*.dwg"))
drawing = ezdxf.new("R2010")
drawing.modelspace().add_mtext("Converted DWG annotation")
drawing.saveas(output_dir / source_name.with_suffix(".dxf").name)
""",
        encoding="utf-8",
    )
    path.chmod(0o700)


def make_client(root: Path, max_parse_bytes: int = 1_048_576) -> TestClient:
    preview_root = root / "previews"
    preview_root.mkdir(exist_ok=True)
    return TestClient(
        create_app(
            Settings(
                PARSER_INTERNAL_TOKEN=TOKEN,
                RAW_DOCS_PATH=root,
                MAX_PARSE_BYTES=max_parse_bytes,
                PREVIEW_ARTIFACTS_PATH=preview_root,
            )
        )
    )


def payload(path: Path) -> dict[str, str]:
    return {
        "jobId": str(uuid4()),
        "documentId": str(uuid4()),
        "storagePath": str(path),
        "mimeType": "text/plain",
    }


def test_parse_requires_internal_token_and_returns_contract(tmp_path: Path) -> None:
    document = tmp_path / "document.txt"
    document.write_text("付款周期为 30 天", encoding="utf-8")
    client = make_client(tmp_path)

    assert client.post("/internal/v1/parse", json=payload(document)).status_code == 401
    response = client.post(
        "/internal/v1/parse",
        headers={"x-internal-token": TOKEN},
        json=payload(document),
    )

    assert response.status_code == 200, response.text
    assert response.json()["elements"][0]["text"] == "付款周期为 30 天"
    assert response.json()["parserVersion"] == "1.1.0"


def test_markdown_preserves_heading_path(tmp_path: Path) -> None:
    document = tmp_path / "guide.md"
    document.write_text("# 财务制度\n\n## 付款流程\n\n验收后 30 天付款", encoding="utf-8")
    client = make_client(tmp_path)
    body = payload(document)
    body["mimeType"] = "text/markdown"

    response = client.post("/internal/v1/parse", headers={"x-internal-token": TOKEN}, json=body)

    assert response.status_code == 200
    assert response.json()["parser"] == "markdown"
    assert response.json()["elements"][-1]["sectionPath"] == ["财务制度", "付款流程"]


def test_docx_preserves_heading_and_table(tmp_path: Path) -> None:
    path = tmp_path / "policy.docx"
    document = Document()
    document.add_heading("制度", level=1)
    document.add_paragraph("正文")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "周期"
    document.save(path)
    client = make_client(tmp_path)
    body = payload(path)
    body["mimeType"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    response = client.post("/internal/v1/parse", headers={"x-internal-token": TOKEN}, json=body)

    assert response.status_code == 200
    assert response.json()["parser"] == "python-docx"
    assert response.json()["elements"][1]["sectionPath"] == ["制度"]
    assert response.json()["elements"][-1]["elementType"] == "table_row"


def test_docx_returns_generated_preview_manifest(
    tmp_path: Path, monkeypatch: MonkeyPatch
) -> None:
    path = tmp_path / "policy.docx"
    document = Document()
    document.add_paragraph("正文")
    document.save(path)
    body = payload(path)
    body["mimeType"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    document_id = body["documentId"]

    def fake_preview(*_: object, **__: object) -> PreviewArtifact:
        return PreviewArtifact(
            storage_key=f"{document_id}.pdf",
            kind="pdf",
            mime_type="application/pdf",
            size_bytes=1024,
            renderer="libreoffice",
            renderer_version="test",
        )

    monkeypatch.setattr(main_module, "generate_office_pdf", fake_preview)
    response = make_client(tmp_path).post(
        "/internal/v1/parse", headers={"x-internal-token": TOKEN}, json=body
    )

    assert response.status_code == 200
    assert response.json()["preview"] == {
        "storageKey": f"{document_id}.pdf",
        "kind": "pdf",
        "mimeType": "application/pdf",
        "sizeBytes": 1024,
        "renderer": "libreoffice",
        "rendererVersion": "test",
    }


def test_xlsx_preserves_sheet_and_header(tmp_path: Path) -> None:
    path = tmp_path / "records.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "付款"
    sheet.append(["项目", "周期"])
    sheet.append(["A", 30])
    workbook.save(path)
    client = make_client(tmp_path)
    body = payload(path)
    body["mimeType"] = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    response = client.post("/internal/v1/parse", headers={"x-internal-token": TOKEN}, json=body)

    assert response.status_code == 200
    assert response.json()["parser"] == "openpyxl"
    assert response.json()["elements"][1]["sheet"] == "付款"
    assert response.json()["elements"][1]["metadata"]["headers"] == ["项目", "周期"]


def test_pdf_routes_to_unstructured_and_preserves_page(
    tmp_path: Path, monkeypatch: MonkeyPatch
) -> None:
    path = tmp_path / "policy.pdf"
    path.write_bytes(b"%PDF-1.4\n%%EOF\n")

    def fake_parse_pdf(
        source: Path, languages: list[str], max_pages: int, max_elements: int
    ) -> tuple[list[ParsedElement], list[str], str]:
        assert source == path
        assert languages == ["ch_sim", "en"]
        assert max_pages == 500
        assert max_elements == 100_000
        return (
            [ParsedElement(text="付款周期为 30 天", element_type="paragraph", page=2)],
            [],
            "test",
        )

    monkeypatch.setattr(main_module, "parse_pdf", fake_parse_pdf)
    client = make_client(tmp_path)
    body = payload(path)
    body["mimeType"] = "application/pdf"

    response = client.post("/internal/v1/parse", headers={"x-internal-token": TOKEN}, json=body)

    assert response.status_code == 200
    assert response.json()["parser"] == "unstructured-pdf"
    assert response.json()["elements"][0]["page"] == 2


def test_pdf_uses_tika_fallback_when_unstructured_fails(
    tmp_path: Path, monkeypatch: MonkeyPatch
) -> None:
    path = tmp_path / "legacy.pdf"
    path.write_bytes(b"%PDF-1.4\n%%EOF\n")

    def failing_parse_pdf(
        source: Path, languages: list[str], max_pages: int, max_elements: int
    ) -> tuple[list[ParsedElement], list[str], str]:
        raise RuntimeError("fixture failure")

    def fake_parse_with_tika(
        source: Path,
        mime_type: str,
        base_url: str,
        timeout_seconds: float,
        max_response_bytes: int,
        max_elements: int,
        parser_version: str,
    ) -> tuple[list[ParsedElement], list[str], str]:
        assert source == path
        assert mime_type == "application/pdf"
        assert base_url == "http://tika:9998"
        assert timeout_seconds == 120
        assert max_response_bytes == 52_428_800
        assert max_elements == 100_000
        assert parser_version == "3.2.3.0"
        return (
            [ParsedElement(text="Tika extracted text", element_type="paragraph")],
            ["TIKA_FALLBACK_USED"],
            parser_version,
        )

    monkeypatch.setattr(main_module, "parse_pdf", failing_parse_pdf)
    monkeypatch.setattr(main_module, "parse_with_tika", fake_parse_with_tika)
    client = TestClient(
        create_app(
            Settings(
                PARSER_INTERNAL_TOKEN=TOKEN,
                RAW_DOCS_PATH=tmp_path,
                TIKA_ENABLED=True,
            )
        )
    )
    body = payload(path)
    body["mimeType"] = "application/pdf"

    response = client.post("/internal/v1/parse", headers={"x-internal-token": TOKEN}, json=body)

    assert response.status_code == 200
    assert response.json()["parser"] == "apache-tika"
    assert response.json()["warnings"] == ["TIKA_FALLBACK_USED"]


def test_pdf_security_validation_does_not_fall_back_to_tika(
    tmp_path: Path, monkeypatch: MonkeyPatch
) -> None:
    path = tmp_path / "encrypted.pdf"
    path.write_bytes(b"%PDF-1.4\n%%EOF\n")

    def reject_encrypted_pdf(
        source: Path, languages: list[str], max_pages: int, max_elements: int
    ) -> tuple[list[ParsedElement], list[str], str]:
        raise ValueError("PDF 文件已加密")

    monkeypatch.setattr(main_module, "parse_pdf", reject_encrypted_pdf)
    client = TestClient(
        create_app(
            Settings(
                PARSER_INTERNAL_TOKEN=TOKEN,
                RAW_DOCS_PATH=tmp_path,
                TIKA_ENABLED=True,
            )
        )
    )
    body = payload(path)
    body["mimeType"] = "application/pdf"

    response = client.post("/internal/v1/parse", headers={"x-internal-token": TOKEN}, json=body)

    assert response.status_code == 422
    assert response.headers["x-parser-error-code"] == "PDF_ENCRYPTED"


def test_image_routes_to_easyocr_and_returns_confidence_warning(
    tmp_path: Path, monkeypatch: MonkeyPatch
) -> None:
    path = tmp_path / "scan.png"
    path.write_bytes(b"\x89PNG\r\n\x1a\nfixture")
    settings = Settings(PARSER_INTERNAL_TOKEN=TOKEN, RAW_DOCS_PATH=tmp_path)
    expected_model_storage_path = settings.ocr_model_storage_path
    expected_user_network_directory = settings.parser_temp_path / "easyocr-user-network"

    def fake_parse_image(
        source: Path,
        model_path: Path,
        user_network_directory: Path,
        languages: list[str],
        max_pixels: int,
        max_elements: int,
        threshold: float,
    ) -> tuple[list[ParsedElement], list[str], str]:
        assert source == path
        assert model_path == expected_model_storage_path
        assert user_network_directory == expected_user_network_directory
        assert languages == ["ch_sim", "en"]
        assert max_pixels == 40_000_000
        assert max_elements == 100_000
        assert threshold == 0.5
        return (
            [
                ParsedElement(
                    text="低置信度文字",
                    element_type="ocr_text",
                    page=1,
                    bbox=[0, 0, 20, 10],
                    metadata={"confidence": 0.42},
                )
            ],
            ["OCR_LOW_CONFIDENCE_ELEMENTS:1"],
            "1.7.2",
        )

    monkeypatch.setattr(main_module, "parse_image", fake_parse_image)
    client = TestClient(create_app(settings))
    body = payload(path)
    body["mimeType"] = "image/png"

    response = client.post("/internal/v1/parse", headers={"x-internal-token": TOKEN}, json=body)

    assert response.status_code == 200
    assert response.json()["parser"] == "easyocr"
    assert response.json()["elements"][0]["metadata"]["confidence"] == 0.42
    assert response.json()["warnings"] == ["OCR_LOW_CONFIDENCE_ELEMENTS:1"]


def test_image_uses_controlled_tmpfs_for_easyocr_user_network(
    tmp_path: Path, monkeypatch: MonkeyPatch
) -> None:
    image_path = tmp_path / "scan.png"
    image_path.write_bytes(b"fixture")
    reader_inputs: dict[str, object] = {}
    user_network_directory = (
        Settings(PARSER_INTERNAL_TOKEN=TOKEN, RAW_DOCS_PATH=tmp_path).parser_temp_path
        / "easyocr-user-network"
    )

    class FakeImage:
        size = (20, 10)

        def __enter__(self) -> "FakeImage":
            return self

        def __exit__(self, *_: object) -> None:
            return None

        def verify(self) -> None:
            return None

    class FakeReader:
        def __init__(self, languages: list[str], **kwargs: object) -> None:
            reader_inputs["languages"] = languages
            reader_inputs.update(kwargs)

        def readtext(self, *_: object, **__: object) -> list[object]:
            return [[[[0, 0], [20, 0], [20, 10], [0, 10]], "OCR text", 0.9]]

    class FakeEasyOcr:
        __version__ = "test"
        Reader = FakeReader

    def fake_import_module(name: str) -> object:
        if name == "PIL.Image":
            return type("FakePillow", (), {"open": staticmethod(lambda _: FakeImage())})
        if name == "easyocr":
            return FakeEasyOcr
        raise AssertionError(f"unexpected import: {name}")

    monkeypatch.setattr(image_parser, "import_module", fake_import_module)

    elements, warnings, version = image_parser.parse_image(
        image_path,
        Path("/opt/easyocr-models"),
        user_network_directory,
        ["ch_sim", "en"],
        40_000_000,
        100_000,
        0.5,
    )

    assert len(elements) == 1
    assert warnings == []
    assert version == "test"
    assert reader_inputs["languages"] == ["ch_sim", "en"]
    assert reader_inputs["model_storage_directory"] == "/opt/easyocr-models"
    assert reader_inputs["user_network_directory"] == str(user_network_directory)
    assert reader_inputs["download_enabled"] is False


def test_dxf_extracts_summary_text_block_attributes_and_dimensions(tmp_path: Path) -> None:
    path = tmp_path / "drawing.dxf"
    drawing = ezdxf.new("R2010")
    drawing.layers.add("ANNOTATION")
    drawing.layers.add("DIMENSIONS")
    block = drawing.blocks.new("TITLE_BLOCK")
    block.add_attdef("PROJECT", (0, 0), text="未填写", dxfattribs={"layer": "ANNOTATION"})
    insert = drawing.modelspace().add_blockref("TITLE_BLOCK", (0, 0))
    insert.add_auto_attribs({"PROJECT": "NexusKB CAD"})
    drawing.modelspace().add_mtext("Payment term: 30 days", dxfattribs={"layer": "ANNOTATION"})
    dimension = drawing.modelspace().add_linear_dim(
        base=(0, 3),
        p1=(0, 0),
        p2=(12.5, 0),
        dxfattribs={"layer": "DIMENSIONS"},
    )
    dimension.render()
    drawing.saveas(path)
    client = make_client(tmp_path)
    body = payload(path)
    body["mimeType"] = "image/vnd.dxf"

    response = client.post("/internal/v1/parse", headers={"x-internal-token": TOKEN}, json=body)

    assert response.status_code == 200
    parsed = response.json()
    assert parsed["parser"] == "ezdxf"
    assert parsed["parserVersion"] == ezdxf.__version__
    assert parsed["elements"][0]["elementType"] == "cad_summary"
    assert "ANNOTATION" in parsed["elements"][0]["metadata"]["layers"]
    assert any(element["text"] == "NexusKB CAD" for element in parsed["elements"])
    assert any(element["text"] == "Payment term: 30 days" for element in parsed["elements"])
    assert any(element["elementType"] == "cad_dimension" for element in parsed["elements"])
    assert parsed["preview"]["kind"] == "svg"
    assert (tmp_path / "previews" / parsed["preview"]["storageKey"]).is_file()


def test_dxf_rejects_entity_limit(tmp_path: Path) -> None:
    path = tmp_path / "large.dxf"
    drawing = ezdxf.new("R2010")
    drawing.modelspace().add_text("one")
    drawing.modelspace().add_text("two")
    drawing.saveas(path)
    client = TestClient(
        create_app(
            Settings(
                PARSER_INTERNAL_TOKEN=TOKEN,
                RAW_DOCS_PATH=tmp_path,
                MAX_CAD_ENTITIES=1,
            )
        )
    )
    body = payload(path)
    body["mimeType"] = "application/dxf"

    response = client.post("/internal/v1/parse", headers={"x-internal-token": TOKEN}, json=body)

    assert response.status_code == 422
    assert response.json()["detail"] == "CAD 实体数量超过限制"
    assert response.headers["x-parser-error-code"] == "CAD_ENTITY_LIMIT_EXCEEDED"


def test_dxf_rejects_corrupted_content_without_leaking_path(tmp_path: Path) -> None:
    path = tmp_path / "corrupted.dxf"
    path.write_bytes(b"0\nSECTION\n2\nENTITIES\nnot-valid-dxf")
    client = make_client(tmp_path)
    body = payload(path)
    body["mimeType"] = "image/vnd.dxf"

    response = client.post("/internal/v1/parse", headers={"x-internal-token": TOKEN}, json=body)

    assert response.status_code == 422
    assert response.json()["detail"] == "DXF 文件损坏或格式不受支持"
    assert str(tmp_path) not in response.text


def test_dwg_is_converted_to_dxf_and_parsed(tmp_path: Path) -> None:
    path = tmp_path / "drawing.dwg"
    path.write_bytes(b"AC1032" + b"\0" * 64)
    converter = tmp_path / "fake-oda-file-converter"
    make_fake_dwg_converter(converter)
    client = TestClient(
        create_app(
            Settings(
                PARSER_INTERNAL_TOKEN=TOKEN,
                RAW_DOCS_PATH=tmp_path,
                DWG_CONVERSION_ENABLED=True,
                DWG_CONVERTER_EXECUTABLE=converter,
                DWG_CONVERTER_RELEASE="test",
                PARSER_TEMP_PATH=tmp_path,
            )
        )
    )
    body = payload(path)
    body["mimeType"] = "image/vnd.dwg"

    response = client.post("/internal/v1/parse", headers={"x-internal-token": TOKEN}, json=body)

    assert response.status_code == 200, response.text
    parsed = response.json()
    assert parsed["parser"] == "oda-file-converter+ezdxf"
    assert parsed["parserVersion"].startswith("oda-test+ezdxf-")
    assert "DWG_CONVERTED_TO_DXF" in parsed["warnings"]
    assert "DWG_SOURCE_VERSION:AC1032" in parsed["warnings"]
    assert any(element["text"] == "Converted DWG annotation" for element in parsed["elements"])


def test_dwg_rejects_unavailable_converter_and_forged_signature(tmp_path: Path) -> None:
    path = tmp_path / "drawing.dwg"
    path.write_bytes(b"AC1032" + b"\0" * 64)
    body = payload(path)
    body["mimeType"] = "image/vnd.dwg"

    unavailable_client = TestClient(
        create_app(
            Settings(
                PARSER_INTERNAL_TOKEN=TOKEN,
                RAW_DOCS_PATH=tmp_path,
                DWG_CONVERSION_ENABLED=True,
                DWG_CONVERTER_EXECUTABLE=tmp_path / "missing-oda-file-converter",
                PARSER_TEMP_PATH=tmp_path,
            )
        )
    )
    unavailable = unavailable_client.post(
        "/internal/v1/parse", headers={"x-internal-token": TOKEN}, json=body
    )
    assert unavailable.status_code == 503
    assert unavailable.json()["detail"] == "DWG 转换器未就绪"

    converter = tmp_path / "fake-oda-file-converter"
    make_fake_dwg_converter(converter)
    path.write_bytes(b"NOTDWG")
    enabled_client = TestClient(
        create_app(
            Settings(
                PARSER_INTERNAL_TOKEN=TOKEN,
                RAW_DOCS_PATH=tmp_path,
                DWG_CONVERSION_ENABLED=True,
                DWG_CONVERTER_EXECUTABLE=converter,
                PARSER_TEMP_PATH=tmp_path,
            )
        )
    )
    invalid = enabled_client.post(
        "/internal/v1/parse", headers={"x-internal-token": TOKEN}, json=body
    )
    assert invalid.status_code == 422
    assert invalid.json()["detail"] == "DWG 版本不受支持或文件签名无效"
    assert invalid.headers["x-parser-error-code"] == "DWG_VERSION_UNSUPPORTED"


def test_dwg_rejects_disabled_conversion(tmp_path: Path) -> None:
    path = tmp_path / "drawing.dwg"
    path.write_bytes(b"AC1032" + b"\0" * 64)
    body = payload(path)
    body["mimeType"] = "image/vnd.dwg"
    disabled_client = TestClient(
        create_app(
            Settings(
                PARSER_INTERNAL_TOKEN=TOKEN,
                RAW_DOCS_PATH=tmp_path,
                DWG_CONVERSION_ENABLED=False,
                PARSER_TEMP_PATH=tmp_path,
            )
        )
    )

    disabled = disabled_client.post(
        "/internal/v1/parse", headers={"x-internal-token": TOKEN}, json=body
    )

    assert disabled.status_code == 503
    assert disabled.json()["detail"] == "DWG 格式转换未启用"


def test_parse_rejects_path_outside_root_and_symlink(tmp_path: Path) -> None:
    root = tmp_path / "root"
    root.mkdir()
    outside = tmp_path / "outside.txt"
    outside.write_text("secret", encoding="utf-8")
    client = make_client(root)

    outside_response = client.post(
        "/internal/v1/parse",
        headers={"x-internal-token": TOKEN},
        json=payload(outside),
    )
    assert outside_response.status_code == 422

    link = root / "link.txt"
    link.symlink_to(outside)
    link_response = client.post(
        "/internal/v1/parse",
        headers={"x-internal-token": TOKEN},
        json=payload(link),
    )
    assert link_response.status_code == 422


def test_parse_rejects_empty_and_oversized_files(tmp_path: Path) -> None:
    client = make_client(tmp_path, max_parse_bytes=1024)
    empty = tmp_path / "empty.txt"
    empty.write_text("", encoding="utf-8")
    assert (
        client.post(
            "/internal/v1/parse", headers={"x-internal-token": TOKEN}, json=payload(empty)
        ).status_code
        == 422
    )

    oversized = tmp_path / "large.txt"
    oversized.write_text("x" * 1025, encoding="utf-8")
    assert (
        client.post(
            "/internal/v1/parse", headers={"x-internal-token": TOKEN}, json=payload(oversized)
        ).status_code
        == 422
    )
