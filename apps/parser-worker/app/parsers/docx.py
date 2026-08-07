from pathlib import Path

from docx import Document

from app.schemas import ParsedElement


def parse_docx(path: Path, max_elements: int) -> list[ParsedElement]:
    document = Document(str(path))
    elements: list[ParsedElement] = []
    section_path: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading "):
            try:
                level = max(1, min(6, int(style_name.removeprefix("Heading "))))
            except ValueError:
                level = 1
            section_path = section_path[: level - 1] + [text]
            element_type = "heading"
        else:
            element_type = "paragraph"
        elements.append(
            ParsedElement(text=text, elementType=element_type, sectionPath=section_path.copy())
        )
        _check_limit(elements, max_elements)
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                elements.append(
                    ParsedElement(
                        text="\t".join(values),
                        elementType="table_row",
                        sectionPath=section_path.copy(),
                        metadata={"tableIndex": table_index, "rowIndex": row_index},
                    )
                )
                _check_limit(elements, max_elements)
    return elements


def _check_limit(elements: list[ParsedElement], max_elements: int) -> None:
    if len(elements) > max_elements:
        raise ValueError("解析结果元素数量超过限制")
